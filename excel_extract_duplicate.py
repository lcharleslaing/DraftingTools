import openpyxl
import json
import os
from pathlib import Path
from datetime import datetime
import sys

def extract_excel_data(file_path):
    """Extract all data from an Excel file including formulas, formatting, and styling"""
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=False)
        data = {
            "file_path": file_path,
            "extraction_date": datetime.now().isoformat(),
            "sheets": {}
        }
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = {
                "name": sheet_name,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "cells": {}
            }
            
            # Extract all non-empty cells
            for row in range(1, sheet.max_row + 1):
                for col in range(1, sheet.max_column + 1):
                    cell = sheet.cell(row=row, column=col)
                    if cell.value is not None:
                        cell_ref = f"{cell.coordinate}"
                        cell_data = {
                            "value": str(cell.value),
                            "data_type": cell.data_type,
                            "number_format": cell.number_format
                        }
                        
                        # Add formula if it exists
                        if cell.data_type == 'f':
                            cell_data["formula"] = str(cell.value)
                        else:
                            cell_data["formula"] = None
                        
                        # Add font information
                        cell_data["font"] = {
                            "name": cell.font.name,
                            "size": cell.font.size,
                            "bold": cell.font.bold,
                            "italic": cell.font.italic,
                            "color": str(cell.font.color.rgb) if cell.font.color and cell.font.color.rgb else None
                        }
                        
                        # Add alignment information
                        cell_data["alignment"] = {
                            "horizontal": cell.alignment.horizontal,
                            "vertical": cell.alignment.vertical,
                            "wrap_text": cell.alignment.wrap_text,
                            "indent": cell.alignment.indent
                        }
                        
                        # Add border information
                        cell_data["border"] = {
                            "left": str(cell.border.left.style) if cell.border.left.style else None,
                            "right": str(cell.border.right.style) if cell.border.right.style else None,
                            "top": str(cell.border.top.style) if cell.border.top.style else None,
                            "bottom": str(cell.border.bottom.style) if cell.border.bottom.style else None
                        }
                        
                        # Add fill information
                        cell_data["fill"] = {
                            "fgColor": str(cell.fill.fgColor.rgb) if cell.fill.fgColor and cell.fill.fgColor.rgb else None,
                            "bgColor": str(cell.fill.bgColor.rgb) if cell.fill.bgColor and cell.fill.bgColor.rgb else None,
                            "patternType": str(cell.fill.patternType) if cell.fill.patternType else None
                        }
                        
                        # Add protection information
                        cell_data["protection"] = {
                            "locked": cell.protection.locked,
                            "hidden": cell.protection.hidden
                        }
                        
                        sheet_data["cells"][cell_ref] = cell_data
            
            # Add sheet properties
            sheet_data["properties"] = {
                "sheet_format": {
                    "defaultRowHeight": sheet.sheet_format.defaultRowHeight,
                    "defaultColWidth": sheet.sheet_format.defaultColWidth
                },
                "sheet_view": {
                    "showGridLines": sheet.sheet_view.showGridLines,
                    "showRowColHeaders": sheet.sheet_view.showRowColHeaders
                }
            }
            
            data["sheets"][sheet_name] = sheet_data
        
        return data
        
    except Exception as e:
        return {"error": str(e), "file_path": file_path, "extraction_date": datetime.now().isoformat()}

def extract_word_data(file_path):
    """Extract data from Word documents (requires python-docx)"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        data = {
            "file_path": file_path,
            "extraction_date": datetime.now().isoformat(),
            "paragraphs": [],
            "tables": []
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                para_data = {
                    "text": para.text,
                    "style": para.style.name,
                    "runs": []
                }
                
                # Extract runs (formatted text segments)
                for run in para.runs:
                    run_data = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None,
                        "font_color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None
                    }
                    para_data["runs"].append(run_data)
                
                data["paragraphs"].append(para_data)
        
        # Extract tables
        for table in doc.tables:
            table_data = {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "cells": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_data = {
                        "text": cell.text,
                        "paragraphs": []
                    }
                    
                    # Extract cell paragraphs
                    for para in cell.paragraphs:
                        if para.text.strip():
                            para_data = {
                                "text": para.text,
                                "style": para.style.name
                            }
                            cell_data["paragraphs"].append(para_data)
                    
                    row_data.append(cell_data)
                table_data["cells"].append(row_data)
            
            data["tables"].append(table_data)
        
        return data
        
    except ImportError:
        return {"error": "python-docx not installed", "file_path": file_path}
    except Exception as e:
        return {"error": str(e), "file_path": file_path, "extraction_date": datetime.now().isoformat()}

def extract_pdf_data(file_path):
    """Extract data from PDF files (requires PyPDF2)"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            data = {
                "file_path": file_path,
                "extraction_date": datetime.now().isoformat(),
                "pages": []
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_data = {
                    "page_number": page_num + 1,
                    "text": page.extract_text(),
                    "mediabox": {
                        "width": float(page.mediabox.width),
                        "height": float(page.mediabox.height)
                    }
                }
                data["pages"].append(page_data)
            
            return data
            
    except ImportError:
        return {"error": "PyPDF2 not installed", "file_path": file_path}
    except Exception as e:
        return {"error": str(e), "file_path": file_path, "extraction_date": datetime.now().isoformat()}

def process_files(directory, file_types=None):
    """Process files in a directory and extract their data"""
    if file_types is None:
        file_types = ['.xlsx', '.xls', '.docx', '.doc', '.pdf']
    
    all_data = {
        "extraction_date": datetime.now().isoformat(),
        "directory": directory,
        "excel_files": [],
        "word_files": [],
        "pdf_files": [],
        "other_files": []
    }
    
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            print(f"Processing: {file_path}")
            
            if file_ext in ['.xlsx', '.xls']:
                data = extract_excel_data(file_path)
                all_data["excel_files"].append(data)
            elif file_ext in ['.docx', '.doc']:
                data = extract_word_data(file_path)
                all_data["word_files"].append(data)
            elif file_ext == '.pdf':
                data = extract_pdf_data(file_path)
                all_data["pdf_files"].append(data)
            else:
                # For other files, just record basic info
                try:
                    stat = os.stat(file_path)
                    file_data = {
                        "file_path": file_path,
                        "file_name": file,
                        "file_size": stat.st_size,
                        "created_date": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        "modified_date": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "file_type": file_ext
                    }
                    all_data["other_files"].append(file_data)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
    
    return all_data

def main():
    """Main function to run the extraction"""
    if len(sys.argv) > 1:
        directory = sys.argv[1]
    else:
        directory = os.getcwd()
    
    print(f"Processing files in: {directory}")
    print("Supported file types: Excel (.xlsx, .xls), Word (.docx, .doc), PDF (.pdf)")
    print("=" * 50)
    
    # Process files
    data = process_files(directory)
    
    # Save to JSON
    output_file = "file_extraction_data.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print("=" * 50)
    print(f"Extraction complete. Data saved to: {output_file}")
    print(f"Excel files processed: {len(data['excel_files'])}")
    print(f"Word files processed: {len(data['word_files'])}")
    print(f"PDF files processed: {len(data['pdf_files'])}")
    print(f"Other files processed: {len(data['other_files'])}")
    
    # Print summary of errors
    errors = []
    for file_list in [data['excel_files'], data['word_files'], data['pdf_files']]:
        for file_data in file_list:
            if 'error' in file_data:
                errors.append(f"{file_data['file_path']}: {file_data['error']}")
    
    if errors:
        print("\nErrors encountered:")
        for error in errors:
            print(f"  - {error}")

if __name__ == "__main__":
    main()
